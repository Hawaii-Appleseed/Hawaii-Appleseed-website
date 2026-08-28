#!/usr/bin/env python3
"""
Render Hawaiʻi Appleseed testimony to .docx (for submission) and .html
(for the clipboard -> Google Doc route in ~/.claude/drive-routes.yml).

Input is the plain-text shape the appleseed-testimony skill produces:

    Testimony of the Hawaiʻi Appleseed Center for Law and Economic Justice
    Support for HB 1884 – Relating to Pedestrian Safety
    House Committee on Transportation
    Thursday, February 12, 2026, at 2:00 PM

    Dear Chair Smith, Vice Chair Jones, and Members of the Committee:

    <body paragraphs, one per line>

    Mahalo for the opportunity to testify.

    Hawaiʻi Appleseed Center for Law and Economic Justice

    ________________

    [1] Source, "Title," Publisher, Date. https://...

Usage:
    python render_testimony.py draft.txt                 # -> draft.docx + draft.html
    python render_testimony.py draft.txt -o out/HB1884   # -> out/HB1884.docx + .html
    python render_testimony.py draft.txt --docx-only
"""
from __future__ import annotations

import argparse, base64, html, os, re, sys
from pathlib import Path

ASSETS = Path(__file__).resolve().parent / "assets"
LOGO = ASSETS / "appleseed-horizontal-green.png"

ORG = "Hawaiʻi Appleseed Center for Law and Economic Justice"
BOILERPLATE = (
    "The Hawaiʻi Appleseed Center for Law and Economic Justice advocates for economic "
    "justice for and with Hawaiʻi’s people. We envision a Hawaiʻi that puts its people "
    "first—where everyone can meet their basic needs while living happy, healthy and "
    "creative lives."
)
FONT, SIZE_PT = "Arial", 11
RULE = "_" * 16


# ---------------------------------------------------------------- parsing
def parse(text: str) -> dict:
    lines = [l.rstrip() for l in text.replace("﻿", "").split("\n")]
    lines = [l for l in lines if l.strip()]

    gi = next((i for i, l in enumerate(lines) if re.match(r"^Dear\b", l)), None)
    if gi is None:
        sys.exit("error: no 'Dear Chair ...' greeting line found")

    header = lines[:gi]
    greeting = lines[gi]
    rest = lines[gi + 1 :]

    # footnotes live after the last rule
    ri = next((i for i, l in enumerate(rest) if re.match(r"^_{6,}$", l.strip())), len(rest))
    notes = [l for l in rest[ri:] if re.match(r"^\[\d+\]", l.strip())]
    upper = rest[:ri]

    # strip a trailing signature line and pull out the closing
    sig_i = next(
        (i for i, l in enumerate(upper) if l.strip().rstrip(",") == ORG or
         re.match(r"^Hawai.i Appleseed Center", l.strip())),
        None,
    )
    if sig_i is not None:
        upper = upper[:sig_i]

    ci = next(
        (i for i, l in enumerate(upper)
         if re.match(r"^(Mahalo|Thank you for (the opportunity|your)|Sincerely|Respectfully)", l)
         and i >= len(upper) - 3),
        len(upper),
    )
    body, closing = upper[:ci], upper[ci:]

    # drop the boilerplate if the author already pasted it in
    body = [p for p in body if "advocates for economic justice for and with" not in p]

    return dict(header=header, greeting=greeting, body=body, closing=closing, notes=notes)


# ---------------------------------------------------------------- docx
def _field(paragraph, instr: str):
    """Insert a Word field code (PAGE / NUMPAGES) — python-docx has no API for these."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText"); instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr_el, end):
        run._r.append(el)
    return run


def build_docx(doc_parts: dict, out: Path):
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    d = docx.Document()

    style = d.styles["Normal"]
    style.font.name, style.font.size = FONT, Pt(SIZE_PT)
    style.paragraph_format.space_after = Pt(10)
    style.paragraph_format.line_spacing = 1.15

    sec = d.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(1)
    sec.left_margin = sec.right_margin = Inches(1)

    # page 1 header: the logo. later pages: org name + "Page X of Y".
    sec.different_first_page_header_footer = True
    if LOGO.exists():
        p = sec.first_page_header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run().add_picture(str(LOGO), width=Inches(2.9))
    else:
        sec.first_page_header.paragraphs[0].text = ORG

    run_hdr = sec.header.paragraphs[0]
    run_hdr.text = ORG + "\t\t"
    run_hdr.add_run("Page ")
    _field(run_hdr, " PAGE ")
    run_hdr.add_run(" of ")
    _field(run_hdr, " NUMPAGES ")
    for r in run_hdr.runs:
        r.font.size, r.font.name = Pt(9), FONT
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    def para(text, *, bold=False, size=SIZE_PT, space_after=10, color=None):
        p = d.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        r = p.add_run(text)
        r.bold, r.font.name, r.font.size = bold, FONT, Pt(size)
        if color:
            r.font.color.rgb = color
        return p

    for i, line in enumerate(doc_parts["header"]):
        para(line, bold=(i == 0), space_after=(10 if i == len(doc_parts["header"]) - 1 else 2))
    para(doc_parts["greeting"])
    for p in doc_parts["body"]:
        para(p)
    for p in doc_parts["closing"]:
        para(p)
    para(ORG)

    if doc_parts["notes"]:
        para(RULE, space_after=6)
        para(BOILERPLATE, size=9, space_after=10, color=RGBColor(0x44, 0x44, 0x44))
        para(RULE, space_after=6)
        for n in doc_parts["notes"]:
            para(n, size=9, space_after=4, color=RGBColor(0x44, 0x44, 0x44))

    out.parent.mkdir(parents=True, exist_ok=True)
    d.save(str(out))
    return out


# ---------------------------------------------------------------- html
def build_html(doc_parts: dict, out: Path):
    e = html.escape
    logo = ""
    if LOGO.exists():
        b64 = base64.b64encode(LOGO.read_bytes()).decode()
        logo = (f'<p><img src="data:image/png;base64,{b64}" '
                f'alt="Hawaiʻi Appleseed" style="width:290px;height:auto"></p>')

    parts = [
        '<meta charset="utf-8">',
        f'<div style="font-family:{FONT},sans-serif;font-size:{SIZE_PT}pt;line-height:1.4">',
        logo,
    ]
    for i, line in enumerate(doc_parts["header"]):
        parts.append(f'<p style="margin:0"><b>{e(line)}</b></p>' if i == 0
                     else f'<p style="margin:0">{e(line)}</p>')
    parts.append("<p></p>")
    parts.append(f"<p>{e(doc_parts['greeting'])}</p>")
    for p in doc_parts["body"] + doc_parts["closing"]:
        parts.append(f"<p>{e(p)}</p>")
    parts.append(f"<p>{e(ORG)}</p>")
    if doc_parts["notes"]:
        parts.append("<hr>")
        parts.append(f'<p style="font-size:9pt;color:#444">{e(BOILERPLATE)}</p>')
        parts.append("<hr>")
        for n in doc_parts["notes"]:
            parts.append(f'<p style="font-size:9pt;color:#444;margin:0 0 4px">{e(n)}</p>')
    parts.append("</div>")

    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text("\n".join(parts), encoding="utf-8")
    return out


def main():
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("source")
    ap.add_argument("-o", "--out", help="output stem (default: alongside the source)")
    ap.add_argument("--docx-only", action="store_true")
    ap.add_argument("--html-only", action="store_true")
    a = ap.parse_args()

    src = Path(a.source).expanduser()
    if not src.exists():
        sys.exit(f"error: {src} not found")
    stem = Path(a.out).expanduser() if a.out else src.with_suffix("")

    parts = parse(src.read_text(encoding="utf-8"))
    if not LOGO.exists():
        print(f"warning: logo missing at {LOGO} — rendering without letterhead", file=sys.stderr)

    if not a.html_only:
        print(build_docx(parts, stem.with_suffix(".docx")))
    if not a.docx_only:
        print(build_html(parts, stem.with_suffix(".html")))


if __name__ == "__main__":
    main()
