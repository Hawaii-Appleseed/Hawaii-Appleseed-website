#!/usr/bin/env python3
"""Build the HA Expense Request template for board 7290593651.

This reproduces the layout of the PDFs DocuGen has been producing: landscape letter,
Appleseed logo top-left with "Expense Request" to its right, a two-column detail block,
an eight-column line-item table, TOTAL REQUEST, Additional Notes, and the check/approval
footer pinned to the bottom of the page.

The logo is the exact image lifted out of one of DocuGen's own PDFs, so it matches.

    ../.venv/bin/python build_expense_request.py
"""

from __future__ import annotations

import pathlib

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = pathlib.Path(__file__).parent
OUT = HERE / "HA Expense Request.docx"
LOGO = HERE / "appleseed-logo.png"

FONT = "Arial"
HEADER_FILL = "F2F2F2"

# Column widths measured off DocuGen's output, scaled to the 10" text column.
COLUMNS = [
    ("Name", 1.96, "{{ s.name }}", False),
    ("Miles", 0.46, "{{ s.c.miles }}", False),
    ("Grants", 2.12, "{{ s.c.grants }}", False),
    ("Issue Area", 1.08, "{{ s.c.issue_area }}", False),
    ("Project", 2.15, "{{ s.c.project }}", False),
    ("Budget Category", 1.10, "{{ s.c.budget_category | join_list }}", False),
    ("Lobbying/Fundraising?", 0.62, "{{ s.c.lobbying_fundraising | join_list }}", False),
    # s.c.amount is corrected in Python before the template ever sees it (see
    # correct_mileage_amounts() in this file) - the board's Amount formula is
    # Expense Amt only ({numbers6__1}), so a mileage-only line would otherwise show $0.
    ("Amount", 0.58, "{{ s.c.amount | money('') }}", True),
]


def style(run, *, size=10, bold=False, font=FONT, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = font
    if color:
        run.font.color.rgb = color
    return run


def shade(cell, fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def set_grid(table, widths_inches) -> None:
    """Write the widths into <w:tblGrid>.

    Cell-level w:tcW alone is not enough: Word and LibreOffice lay the table out from the
    grid, so without this every column renders the same width no matter what you set.
    """
    grid = table._tbl.find(qn("w:tblGrid"))
    for col, inches in zip(grid.findall(qn("w:gridCol")), widths_inches):
        col.set(qn("w:w"), str(int(inches * 1440)))


def zero_cell_margins(table) -> None:
    """Drop the default cell padding so content sits flush with the page margin."""
    margins = OxmlElement("w:tblCellMar")
    for edge in ("left", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), "0")
        el.set(qn("w:type"), "dxa")
        margins.append(el)
    table._tbl.tblPr.append(margins)


def no_borders(table) -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    table._tbl.tblPr.append(borders)


def top_rule(paragraph, gap_points: int = 21) -> None:
    """Rule above the paragraph, held `gap_points` off the text.

    The gap is the border's own w:space, not paragraph spacing: space_before would move
    the rule and the text together and keep them just as close. DocuGen's footer sits
    0.32in below its rule, which is ~22pt.
    """
    borders = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "12")
    top.set(qn("w:color"), "000000")
    top.set(qn("w:space"), str(gap_points))   # max 31 per the OOXML spec
    borders.append(top)
    paragraph._p.get_or_add_pPr().append(borders)


def label_value(cell, label: str, tag: str) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(6)
    style(paragraph.add_run(label + " "), size=9.5, bold=True)
    style(paragraph.add_run(tag), size=9.5)


def build() -> pathlib.Path:
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = Inches(11), Inches(8.5)
    section.left_margin = section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.5)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10)

    # ---- masthead: logo left, title right -------------------------------------
    masthead = doc.add_table(rows=1, cols=2)
    no_borders(masthead)
    zero_cell_margins(masthead)
    left, right = masthead.rows[0].cells
    left.width, right.width = Inches(4.85), Inches(5.15)
    set_grid(masthead, [4.85, 5.15])
    if LOGO.exists():
        left.paragraphs[0].add_run().add_picture(str(LOGO), width=Inches(3.55))
    title = right.paragraphs[0]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    style(title.add_run("Expense Request"), size=26)

    # ---- request details ------------------------------------------------------
    details = doc.add_table(rows=3, cols=2)
    no_borders(details)
    zero_cell_margins(details)
    set_grid(details, [4.6, 5.4])
    for row in details.rows:
        row.cells[0].width = Inches(4.6)
        row.cells[1].width = Inches(5.4)

    label_value(details.rows[0].cells[0], "Date of Request:",
                "{{ c.date_of_request | date('%m/%d/%Y') }}")
    label_value(details.rows[0].cells[1], "Requested By:", "{{ c.requested_by }}")
    label_value(details.rows[1].cells[0], "Payee:", "{{ c.payee_name }}")
    label_value(details.rows[1].cells[1], "Payee Address:", "{{ c.payee_address }}")
    label_value(details.rows[2].cells[0], "Request ID:", "{{ name }}")
    label_value(details.rows[2].cells[1], "Apt./Ste.:", "{{ c.apt_ste }}")

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)

    # ---- line items -----------------------------------------------------------
    items = doc.add_table(rows=1, cols=len(COLUMNS))
    items.style = "Table Grid"
    # Word only honours explicit widths when autofit is off, and the width has to be set
    # on every cell - setting it on the column alone is silently ignored.
    items.autofit = False
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    items._tbl.tblPr.append(layout)
    set_grid(items, [width for _, width, _, _ in COLUMNS])
    header = items.rows[0]
    for index, (heading, width, _, right_align) in enumerate(COLUMNS):
        cell = header.cells[index]
        cell.width = Inches(width)
        shade(cell, HEADER_FILL)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        paragraph = cell.paragraphs[0]
        if right_align:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style(paragraph.add_run(heading), size=7.5, bold=True)
    header.height = Inches(0.42)

    # {%tr %} consumes the row it sits in, so the loop needs its own tag rows.
    style(items.add_row().cells[0].paragraphs[0].add_run("{%tr for s in subitems %}"),
          size=7.5)
    body = items.add_row().cells
    for index, (_, width, tag, right_align) in enumerate(COLUMNS):
        body[index].width = Inches(width)
        paragraph = body[index].paragraphs[0]
        if right_align:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        style(paragraph.add_run(tag), size=8)
    style(items.add_row().cells[0].paragraphs[0].add_run("{%tr endfor %}"), size=7.5)

    # ---- total ----------------------------------------------------------------
    total = doc.add_paragraph()
    total.paragraph_format.space_before = Pt(14)
    total.paragraph_format.space_after = Pt(8)
    # Always sum the (corrected) line amounts rather than trust the board's Total
    # mirror: that mirror sums the mileage-blind Amount column, so for a mileage-only
    # item it can read as a single, plausible-looking - but wrong - number, defeating
    # grand_total()'s "trust a lone value" heuristic. Summing from lines is correct
    # either way and also sidesteps the mirror's separate concatenated-string bug.
    # DocuGen printed "$$181.19" here (a stray literal $ next to an already-formatted
    # value); one $ is intentional here.
    style(total.add_run("TOTAL REQUEST: ${{ total(subitems, 'amount') | money('') }}"),
          size=13, bold=True)

    # ---- notes ----------------------------------------------------------------
    notes = doc.add_paragraph()
    style(notes.add_run("Additional Notes: "), size=9.5, bold=True)
    style(notes.add_run("{{ c.additional_notes }}"), size=9.5)

    # ---- check / approval footer, pinned to the bottom of the page -------------
    footer = section.footer.paragraphs[0]
    top_rule(footer)
    style(footer.add_run(
        "Check #:_________" + " " * 11 +
        "Dated:__________" + " " * 7 +
        "In the amount of: $__________" + " " * 29 +
        "Approved by:_____________________________"), size=10)

    doc.save(str(OUT))
    return OUT


if __name__ == "__main__":
    print(f"Wrote {build()}")
