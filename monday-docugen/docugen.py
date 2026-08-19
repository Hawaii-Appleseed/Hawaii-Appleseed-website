#!/usr/bin/env python3
"""docugen — generate documents from monday.com board data using .docx templates.

A local, scriptable stand-in for the DocuGen monday.com app.

Subcommands:
  columns    List a board's columns (ids, titles, types) so you can write placeholders.
  scaffold   Write a starter .docx template containing every available placeholder.
  generate   Render a template with data from one item (item mode) or many (board mode).

Templates are ordinary Word documents containing Jinja2 tags, e.g.
    Dear {{ c.client_name }},
    Your quote total is {{ c.total | money }}, due {{ c.due_date | date('%B %-d, %Y') }}.
"""

from __future__ import annotations

import argparse
import datetime as dt
import json
import os
import pathlib
import re
import shutil
import subprocess
import sys
import tempfile

import requests

API_URL = "https://api.monday.com/v2"
FILE_URL = "https://api.monday.com/v2/file"
API_VERSION = "2024-10"
TOKEN_FILE = pathlib.Path.home() / ".monday_api_key"


# --------------------------------------------------------------------------- API


class MondayError(RuntimeError):
    pass


def get_token(explicit: str | None = None) -> str:
    token = explicit or os.environ.get("MONDAY_API_TOKEN")
    if not token and TOKEN_FILE.exists():
        token = TOKEN_FILE.read_text().strip()
    if not token:
        raise MondayError(
            f"No API token. Set MONDAY_API_TOKEN, pass --token, or write it to {TOKEN_FILE}.\n"
            "Get one in monday.com: avatar → Developers → My access tokens."
        )
    return token


def gql(token: str, query: str, variables: dict | None = None) -> dict:
    resp = requests.post(
        API_URL,
        json={"query": query, "variables": variables or {}},
        headers={
            "Authorization": token,
            "API-Version": API_VERSION,
            "Content-Type": "application/json",
        },
        timeout=60,
    )
    if resp.status_code != 200:
        raise MondayError(f"HTTP {resp.status_code} from monday API: {resp.text[:500]}")
    try:
        payload = resp.json()
    except ValueError:
        raise MondayError(f"Non-JSON response from monday API: {resp.text[:500]}")
    if payload.get("errors"):
        raise MondayError("monday API error: " + json.dumps(payload["errors"], indent=2))
    if "data" not in payload:
        raise MondayError(f"Unexpected API payload: {json.dumps(payload)[:500]}")
    return payload["data"]


COLUMN_VALUE_FRAGMENT = """
      id
      type
      text
      value
      column { id title type }
      ... on MirrorValue { display_value }
      ... on BoardRelationValue { display_value }
      ... on DependencyValue { display_value }
      ... on FormulaValue { display_value }
"""

ITEM_FIELDS = f"""
    id
    name
    url
    state
    created_at
    updated_at
    group {{ id title }}
    board {{ id name }}
    column_values {{{COLUMN_VALUE_FRAGMENT}}}
"""

Q_ITEM = f"""
query ($ids: [ID!]) {{
  items (ids: $ids) {{
    {ITEM_FIELDS}
    subitems {{
      id
      name
      column_values {{{COLUMN_VALUE_FRAGMENT}}}
    }}
  }}
}}
"""

Q_BOARD_ITEMS = f"""
query ($board: [ID!], $cursor: String, $limit: Int!) {{
  boards (ids: $board) {{
    id
    name
    items_page (limit: $limit, cursor: $cursor) {{
      cursor
      items {{
        {ITEM_FIELDS}
        subitems {{
          id
          name
          column_values {{{COLUMN_VALUE_FRAGMENT}}}
        }}
      }}
    }}
  }}
}}
"""

Q_COLUMNS = """
query ($board: [ID!]) {
  boards (ids: $board) {
    id
    name
    columns { id title type }
    groups { id title }
  }
}
"""


def fetch_item(token: str, item_id: str) -> dict:
    items = gql(token, Q_ITEM, {"ids": [str(item_id)]})["items"]
    if not items:
        raise MondayError(f"Item {item_id} not found (or the token can't see it).")
    return items[0]


def fetch_board_items(token: str, board_id: str, group: str | None = None,
                      limit: int | None = None) -> tuple[dict, list[dict]]:
    cursor, out, board_meta = None, [], {}
    while True:
        data = gql(token, Q_BOARD_ITEMS,
                   {"board": [str(board_id)], "cursor": cursor, "limit": 100})
        boards = data["boards"]
        if not boards:
            raise MondayError(f"Board {board_id} not found (or the token can't see it).")
        board_meta = {"id": boards[0]["id"], "name": boards[0]["name"]}
        page = boards[0]["items_page"]
        out.extend(page["items"])
        cursor = page.get("cursor")
        if not cursor or (limit and len(out) >= limit):
            break
    if group:
        want = group.lower()
        out = [i for i in out
               if (i.get("group") or {}).get("title", "").lower() == want
               or (i.get("group") or {}).get("id", "") == group]
    if limit:
        out = out[:limit]
    return board_meta, out


# ------------------------------------------------------------------- value coercion


def slug(title: str) -> str:
    s = re.sub(r"[^0-9a-zA-Z]+", "_", (title or "").strip().lower()).strip("_")
    if not s:
        s = "col"
    if s[0].isdigit():
        s = "c_" + s
    return s


def coerce(cv: dict):
    """Turn one monday column_value into a friendly Python value."""
    ctype = cv.get("type") or ""
    text = cv.get("text")
    disp = cv.get("display_value")
    raw = cv.get("value")
    parsed = None
    if raw:
        try:
            parsed = json.loads(raw)
        except (ValueError, TypeError):
            parsed = None

    if ctype in ("mirror", "board_relation", "dependency", "formula"):
        return disp or text or ""

    if ctype == "numbers":
        if text in (None, ""):
            return None
        try:
            f = float(text)
            return int(f) if f.is_integer() else f
        except ValueError:
            return text

    if ctype == "checkbox":
        return bool(parsed and parsed.get("checked") in (True, "true"))

    if ctype in ("date", "timeline"):
        if isinstance(parsed, dict):
            if parsed.get("from"):  # timeline
                return {"from": _to_date(parsed.get("from")), "to": _to_date(parsed.get("to")),
                        "text": text or ""}
            if parsed.get("date"):
                d = _to_date(parsed["date"])
                if parsed.get("time"):
                    try:
                        t = dt.datetime.strptime(parsed["time"], "%H:%M:%S").time()
                        return dt.datetime.combine(d, t)
                    except (ValueError, TypeError):
                        pass
                return d
        return _to_date(text) if text else None

    if ctype in ("people", "multiple-person", "person"):
        return text or ""

    if ctype in ("dropdown", "tags"):
        return [p.strip() for p in (text or "").split(",") if p.strip()]

    if ctype == "file":
        if isinstance(parsed, dict):
            return [f.get("name") or f.get("assetId") for f in parsed.get("files", [])]
        return []

    if ctype == "link" and isinstance(parsed, dict):
        return parsed.get("url") or text or ""

    return text if text is not None else ""


def _to_date(s):
    if not s:
        return None
    try:
        return dt.date.fromisoformat(str(s)[:10])
    except ValueError:
        return s


def build_item_context(item: dict) -> dict:
    by_title, by_id, raw = {}, {}, {}
    for cv in item.get("column_values") or []:
        title = (cv.get("column") or {}).get("title") or cv.get("id")
        val = coerce(cv)
        by_title[slug(title)] = val
        by_id[cv["id"]] = val
        raw[cv["id"]] = {"title": title, "type": cv.get("type"),
                         "text": cv.get("text"), "value": cv.get("value")}
    ctx = {
        "id": item.get("id"),
        "name": item.get("name"),
        "url": item.get("url"),
        "state": item.get("state"),
        "created_at": item.get("created_at"),
        "updated_at": item.get("updated_at"),
        "group": (item.get("group") or {}).get("title", ""),
        "board": (item.get("board") or {}).get("name", ""),
        "board_id": (item.get("board") or {}).get("id", ""),
        "c": by_title,
        "col": by_id,
        "raw": raw,
    }
    subs = item.get("subitems") or []
    ctx["subitems"] = [build_item_context(s) for s in subs]
    ctx["subitem_count"] = len(subs)
    return ctx


def base_context() -> dict:
    now = dt.datetime.now()
    return {"now": now, "today": now.date(),
            "today_str": now.strftime("%B %-d, %Y"),
            "generated_at": now.strftime("%Y-%m-%d %H:%M")}


# ---------------------------------------------------------------------- rendering


def missing(v) -> bool:
    """True for None, empty string, or a Jinja undefined (missing column)."""
    import jinja2
    return v is None or isinstance(v, jinja2.Undefined) or v == ""


def money(v, symbol="$", places=2):
    """Format a value as currency. Pass symbol="" for a bare 1,234.56.

    Formula and mirror columns arrive pre-formatted ("$1,181.19"), so fall back to
    as_number rather than echoing the raw string back into the document.
    """
    if missing(v):
        return ""
    try:
        return f"{symbol}{float(v):,.{places}f}"
    except (TypeError, ValueError):
        pass
    if any(ch.isdigit() for ch in str(v)):
        return f"{symbol}{as_number(v):,.{places}f}"
    return str(v)


def date_filter(v, fmt="%B %-d, %Y"):
    if missing(v):
        return ""
    if isinstance(v, str):
        v = _to_date(v) or v
    if isinstance(v, (dt.date, dt.datetime)):
        try:
            return v.strftime(fmt)
        except ValueError:  # platforms without %-d
            return v.strftime(fmt.replace("%-d", "%d").replace("%-m", "%m"))
    return str(v)


def as_number(v) -> float:
    """Parse a number out of a monday value, including pre-formatted "$1,181.19".

    Formula and mirror columns arrive as display strings, so float() alone silently
    contributes 0 and the total comes out wrong.
    """
    if missing(v):
        return 0.0
    if isinstance(v, (int, float)):
        return float(v)
    cleaned = re.sub(r"[^0-9.\-]", "", str(v))
    try:
        return float(cleaned)
    except ValueError:
        return 0.0


def total(items, key):
    s = 0.0
    if missing(items):
        return 0
    for it in items:
        v = (it.get("c") or {}).get(key) if isinstance(it, dict) else None
        s += as_number(v)
    return int(s) if float(s).is_integer() else s


def join_list(v, sep=", "):
    if missing(v):
        return ""
    if isinstance(v, str):
        return v
    try:
        return sep.join(str(x) for x in v)
    except TypeError:
        return str(v)


SINGLE_NUMBER = re.compile(r"^\s*-?\$?\s*-?[\d,]+(?:\.\d+)?\s*$")


def grand_total(value, items, key="amount"):
    """The document's total, trusting the board only when it gives a single number.

    monday mirror columns configured with function:"sum" still hand the API a
    *concatenation* of the underlying values ("18.50, 142.75") via display_value, so a
    multi-line request would otherwise print nonsense where the total belongs. A lone
    value ("$1,181.19") is trustworthy; anything else gets summed from the line items.
    """
    if not missing(value) and SINGLE_NUMBER.match(str(value)):
        return as_number(value)
    return total(items, key)


def register_filters(env):
    env.filters["money"] = money
    env.filters["date"] = date_filter
    env.filters["join_list"] = join_list
    env.filters["yesno"] = lambda v, y="Yes", n="No": n if missing(v) else (y if v else n)
    env.globals["total"] = total
    env.globals["sum_of"] = total
    env.globals["grand_total"] = grand_total
    env.globals["as_number"] = as_number


def render(template: pathlib.Path, context: dict, out: pathlib.Path) -> pathlib.Path:
    from docxtpl import DocxTemplate
    import jinja2

    doc = DocxTemplate(str(template))
    # autoescape is not optional: a value containing & < > is invalid XML raw, and Word
    # silently drops the surrounding text. "Meals & Entertainment" becomes "Meals".
    env = jinja2.Environment(
        undefined=jinja2.ChainableUndefined,
        autoescape=True,
        # An empty numbers column coerces to None; without this it prints as "None".
        finalize=lambda v: "" if v is None else v,
    )
    register_filters(env)
    doc.render(context, env)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


# --------------------------------------------------------------------------- PDF

SOFFICE_CANDIDATES = [
    "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    "/usr/bin/soffice", "/usr/local/bin/soffice", "/opt/homebrew/bin/soffice",
]


def find_soffice() -> str | None:
    return shutil.which("soffice") or next((p for p in SOFFICE_CANDIDATES
                                            if pathlib.Path(p).exists()), None)


def to_pdf(docx_path: pathlib.Path) -> pathlib.Path:
    pdf_path = docx_path.with_suffix(".pdf")
    soffice = find_soffice()
    if soffice:
        with tempfile.TemporaryDirectory() as td:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir", td, str(docx_path)],
                check=True, capture_output=True, timeout=180,
            )
            produced = pathlib.Path(td) / (docx_path.stem + ".pdf")
            if not produced.exists():
                raise MondayError("LibreOffice ran but produced no PDF.")
            shutil.move(str(produced), str(pdf_path))
        return pdf_path
    # Fall back to Microsoft Word via AppleScript, if installed.
    if sys.platform == "darwin" and pathlib.Path("/Applications/Microsoft Word.app").exists():
        script = f'''
        tell application "Microsoft Word"
            set d to open file name POSIX file "{docx_path}"
            save as d file name POSIX file "{pdf_path}" file format format PDF
            close d saving no
        end tell'''
        subprocess.run(["osascript", "-e", script], check=True, capture_output=True, timeout=180)
        return pdf_path
    raise MondayError(
        "PDF conversion needs LibreOffice (or Microsoft Word on macOS).\n"
        "  brew install --cask libreoffice"
    )


# ------------------------------------------------------------------------- upload


def upload_to_column(token: str, item_id: str, column_id: str, path: pathlib.Path) -> dict:
    """Attach a file to a Files column.

    monday's /v2/file endpoint wants the ids inlined in the query and the file sent as a
    form part literally named `variables[file]`. The generic GraphQL-multipart shape
    (operations + map) is rejected here with "Variable $column ... was not provided".
    """
    if not re.fullmatch(r"[A-Za-z0-9_]+", column_id):
        raise MondayError(f"Suspicious column id: {column_id!r}")
    query = (
        "mutation ($file: File!) { add_file_to_column "
        f'(item_id: {int(item_id)}, column_id: "{column_id}", file: $file)'
        " { id name url } }"
    )
    with path.open("rb") as fh:
        resp = requests.post(
            FILE_URL,
            headers={"Authorization": token, "API-Version": API_VERSION},
            data={"query": query},
            files={"variables[file]": (path.name, fh, "application/octet-stream")},
            timeout=180,
        )
    if resp.status_code != 200:
        raise MondayError(f"Upload failed: HTTP {resp.status_code} {resp.text[:400]}")
    payload = resp.json()
    if payload.get("errors"):
        raise MondayError("Upload error: " + json.dumps(payload["errors"], indent=2))
    return payload["data"]["add_file_to_column"]


# ------------------------------------------------------------------------ scaffold


def scaffold_template(board: dict, out: pathlib.Path, mode: str) -> pathlib.Path:
    from docx import Document

    doc = Document()
    doc.add_heading(f"{board['name']} — template", level=1)
    doc.add_paragraph(
        "Replace this boilerplate with your real layout. Every placeholder below is live: "
        "edit the wording around it, delete what you don't need, and keep the double-brace "
        "tags exactly as they appear."
    )

    if mode == "item":
        doc.add_heading("Item fields", level=2)
        for label, tag in [("Item name", "{{ name }}"), ("Group", "{{ group }}"),
                           ("Board", "{{ board }}"), ("Item URL", "{{ url }}"),
                           ("Generated", "{{ today_str }}")]:
            doc.add_paragraph(f"{label}: {tag}")

        doc.add_heading("Columns", level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Column", "Type", "Placeholder"
        for col in board["columns"]:
            row = table.add_row().cells
            row[0].text = col["title"]
            row[1].text = col["type"]
            row[2].text = "{{ c.%s }}" % slug(col["title"])

        doc.add_heading("Subitems", level=2)
        sub = doc.add_table(rows=1, cols=2)
        sub.style = "Table Grid"
        sub.rows[0].cells[0].text = "Subitem"
        sub.rows[0].cells[1].text = "Status"
        sub.add_row().cells[0].text = "{%tr for s in subitems %}"
        body = sub.add_row().cells
        body[0].text = "{{ s.name }}"
        body[1].text = "{{ s.c.status }}"
        sub.add_row().cells[0].text = "{%tr endfor %}"
    else:
        doc.add_heading("Board report", level=2)
        doc.add_paragraph("Board: {{ board }} — {{ item_count }} items as of {{ today_str }}")
        cols = board["columns"][:4]
        table = doc.add_table(rows=1, cols=1 + len(cols))
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Item"
        for i, col in enumerate(cols):
            hdr[i + 1].text = col["title"]
        table.add_row().cells[0].text = "{%tr for it in items %}"
        body = table.add_row().cells
        body[0].text = "{{ it.name }}"
        for i, col in enumerate(cols):
            body[i + 1].text = "{{ it.c.%s }}" % slug(col["title"])
        table.add_row().cells[0].text = "{%tr endfor %}"

    doc.add_paragraph()
    doc.add_paragraph(
        "Filters: {{ c.amount | money }} · {{ c.due_date | date('%B %-d, %Y') }} · "
        "{{ c.tags | join_list }} · {{ c.done | yesno }} · {{ total(items, 'amount') | money }}"
    )
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


# ---------------------------------------------------------------------------- CLI


def cmd_columns(args) -> int:
    token = get_token(args.token)
    boards = gql(token, Q_COLUMNS, {"board": [str(args.board)]})["boards"]
    if not boards:
        raise MondayError(f"Board {args.board} not found.")
    board = boards[0]
    if args.json:
        print(json.dumps(board, indent=2))
        return 0
    print(f"Board {board['id']}: {board['name']}\n")
    print(f"{'PLACEHOLDER':<34} {'TYPE':<18} COLUMN TITLE")
    print("-" * 80)
    for col in board["columns"]:
        print(f"{'{{ c.' + slug(col['title']) + ' }}':<34} {col['type']:<18} {col['title']}")
    print("\nGroups: " + ", ".join(g["title"] for g in board["groups"]))
    return 0


def cmd_scaffold(args) -> int:
    token = get_token(args.token)
    boards = gql(token, Q_COLUMNS, {"board": [str(args.board)]})["boards"]
    if not boards:
        raise MondayError(f"Board {args.board} not found.")
    out = scaffold_template(boards[0], pathlib.Path(args.out).expanduser(), args.mode)
    print(f"Wrote starter template: {out}")
    return 0


def _out_path(pattern: str, ctx: dict) -> pathlib.Path:
    import jinja2
    env = jinja2.Environment(undefined=jinja2.ChainableUndefined)
    register_filters(env)
    name = env.from_string(pattern).render(**ctx)
    name = re.sub(r'[<>:"|?*\x00-\x1f]', "-", name).strip()
    return pathlib.Path(name).expanduser()


def cmd_generate(args) -> int:
    template = pathlib.Path(args.template).expanduser() if args.template else None
    if template and not template.exists():
        raise MondayError(f"Template not found: {template}")
    token = get_token(args.token)
    outputs: list[tuple[pathlib.Path, dict]] = []

    if args.item:
        for item_id in args.item:
            item = fetch_item(token, item_id)
            ctx = base_context() | build_item_context(item)
            outputs.append((_out_path(args.out, ctx), ctx))
    else:
        board_meta, items = fetch_board_items(token, args.board, args.group, args.limit)
        contexts = [build_item_context(i) for i in items]
        if args.per_item:
            for ctx in contexts:
                full = base_context() | ctx
                outputs.append((_out_path(args.out, full), full))
        else:
            ctx = base_context() | {
                "board": board_meta["name"], "board_id": board_meta["id"],
                "group": args.group or "", "items": contexts,
                "item_count": len(contexts), "name": board_meta["name"],
            }
            outputs.append((_out_path(args.out, ctx), ctx))

    if args.dry_run:
        for path, ctx in outputs:
            print(f"=== {path} ===")
            print(json.dumps(ctx, indent=2, default=str))
        return 0

    if not template:
        raise MondayError("--template is required unless you pass --dry-run.")

    for path, ctx in outputs:
        docx_path = render(template, ctx, path if path.suffix == ".docx"
                           else path.with_suffix(".docx"))
        final = to_pdf(docx_path) if args.pdf else docx_path
        if args.pdf and not args.keep_docx and final != docx_path:
            docx_path.unlink(missing_ok=True)
        print(f"Wrote {final}")
        if args.upload_column:
            target = ctx.get("id") or args.upload_item
            if not target:
                raise MondayError("--upload-column needs an item; use item mode or --upload-item.")
            info = upload_to_column(token, target, args.upload_column, final)
            print(f"  uploaded to item {target} column {args.upload_column} "
                  f"(asset {info['id']})")
    return 0


def main(argv=None) -> int:
    p = argparse.ArgumentParser(
        prog="docugen",
        description="Generate documents from monday.com board data using .docx templates.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""examples:
  docugen.py columns --board 1234567890
  docugen.py scaffold --board 1234567890 --out templates/quote.docx
  docugen.py generate --item 987654321 --template templates/quote.docx \\
      --out 'out/Quote - {{ name }}.docx' --pdf --upload-column files
  docugen.py generate --board 1234567890 --group "Q3" \\
      --template templates/report.docx --out out/report.docx --pdf
""",
    )
    p.add_argument("--token", help="monday API token (else MONDAY_API_TOKEN or ~/.monday_api_key)")
    sub = p.add_subparsers(dest="cmd", required=True)

    c = sub.add_parser("columns", help="list a board's columns and their placeholders")
    c.add_argument("--board", required=True)
    c.add_argument("--json", action="store_true")
    c.set_defaults(func=cmd_columns)

    s = sub.add_parser("scaffold", help="write a starter .docx template for a board")
    s.add_argument("--board", required=True)
    s.add_argument("--out", required=True)
    s.add_argument("--mode", choices=["item", "board"], default="item")
    s.set_defaults(func=cmd_scaffold)

    g = sub.add_parser("generate", help="render a template with monday data")
    src = g.add_mutually_exclusive_group(required=True)
    src.add_argument("--item", nargs="+", help="one or more item ids (item mode)")
    src.add_argument("--board", help="board id (board mode)")
    g.add_argument("--template", help="path to the .docx template")
    g.add_argument("--out", default="out/{{ name }}.docx",
                   help="output path; may contain Jinja tags, e.g. 'out/{{ name }}.docx'")
    g.add_argument("--group", help="board mode: only items in this group (id or title)")
    g.add_argument("--limit", type=int, help="board mode: cap the number of items")
    g.add_argument("--per-item", action="store_true",
                   help="board mode: one document per item instead of one combined document")
    g.add_argument("--pdf", action="store_true", help="also convert to PDF")
    g.add_argument("--keep-docx", action="store_true", help="with --pdf, keep the .docx too")
    g.add_argument("--upload-column", help="upload the result to this Files column id")
    g.add_argument("--upload-item", help="item id to upload to (board mode)")
    g.add_argument("--dry-run", action="store_true",
                   help="print the template context as JSON instead of rendering")
    g.set_defaults(func=cmd_generate)

    args = p.parse_args(argv)
    try:
        return args.func(args)
    except MondayError as exc:
        print(f"error: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
