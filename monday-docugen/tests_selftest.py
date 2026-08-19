import json, pathlib, sys
sys.path.insert(0, str(pathlib.Path.home() / "monday-docugen"))
import docugen as D

board = {
    "id": "111", "name": "Client Projects",
    "columns": [
        {"id": "text", "title": "Client Name", "type": "text"},
        {"id": "status", "title": "Status", "type": "status"},
        {"id": "numbers", "title": "Amount", "type": "numbers"},
        {"id": "date4", "title": "Due Date", "type": "date"},
        {"id": "person", "title": "Owner", "type": "people"},
        {"id": "dropdown", "title": "Services", "type": "dropdown"},
        {"id": "check", "title": "Signed?", "type": "checkbox"},
    ],
    "groups": [{"id": "topics", "title": "Q3"}],
}

import tempfile
out = pathlib.Path(tempfile.mkdtemp(prefix="docugen-test-"))
t_item = D.scaffold_template(board, out / "tpl_item.docx", "item")
t_board = D.scaffold_template(board, out / "tpl_board.docx", "board")
print("scaffolded:", t_item.name, t_board.name)


def cv(cid, title, ctype, text, value=None, display=None):
    d = {"id": cid, "type": ctype, "text": text, "value": json.dumps(value) if value is not None else None,
         "column": {"id": cid, "title": title, "type": ctype}}
    if display is not None:
        d["display_value"] = display
    return d


item = {
    "id": "555", "name": "Acme Rebrand", "url": "https://x.monday.com/boards/111/pulses/555",
    "state": "active", "created_at": "2026-01-05T10:00:00Z", "updated_at": "2026-02-01T10:00:00Z",
    "group": {"id": "topics", "title": "Q3"}, "board": {"id": "111", "name": "Client Projects"},
    "column_values": [
        cv("text", "Client Name", "text", "Acme Corp"),
        cv("status", "Status", "status", "Working on it", {"index": 0}),
        cv("numbers", "Amount", "numbers", "12500.5"),
        cv("date4", "Due Date", "date", "2026-09-30", {"date": "2026-09-30"}),
        cv("person", "Owner", "people", "Dana Lee"),
        cv("dropdown", "Services", "dropdown", "Design, Copywriting"),
        cv("check", "Signed?", "checkbox", "v", {"checked": "true"}),
        cv("mir", "Mirror Col", "mirror", None, None, "Mirrored value"),
        cv("empty", "Notes", "long-text", None),
    ],
    "subitems": [
        {"id": "1", "name": "Logo", "column_values": [cv("status", "Status", "status", "Done")]},
        {"id": "2", "name": "Website", "column_values": [cv("status", "Status", "status", "Stuck")]},
    ],
}

ctx = D.base_context() | D.build_item_context(item)
print(json.dumps({k: v for k, v in ctx.items() if k in ("name", "group", "c", "subitem_count")}, indent=2, default=str))
assert ctx["c"]["amount"] == 12500.5, ctx["c"]["amount"]
assert ctx["c"]["signed"] is True
assert ctx["c"]["services"] == ["Design", "Copywriting"]
assert str(ctx["c"]["due_date"]) == "2026-09-30"
assert ctx["c"]["mirror_col"] == "Mirrored value"
assert ctx["c"]["notes"] == ""

r1 = D.render(t_item, ctx, out / "rendered_item.docx")
print("rendered:", r1)

items = [D.build_item_context(item), D.build_item_context({**item, "id": "556", "name": "Beta Launch", "subitems": []})]
bctx = D.base_context() | {"board": "Client Projects", "items": items, "item_count": 2, "name": "Client Projects"}
r2 = D.render(t_board, bctx, out / "rendered_board.docx")
print("rendered:", r2)

# custom template exercising filters + missing keys
from docx import Document
doc = Document()
doc.add_paragraph("Dear {{ c.client_name }} ({{ c.owner }}),")
doc.add_paragraph("Total {{ c.amount | money }} due {{ c.due_date | date('%B %-d, %Y') }}.")
doc.add_paragraph("Services: {{ c.services | join_list }} | Signed: {{ c.signed | yesno }}")
doc.add_paragraph("Missing on purpose: [{{ c.does_not_exist }}] [{{ c.nope.deeper }}]")
doc.add_paragraph("{% for s in subitems %}- {{ s.name }}: {{ s.c.status }}\n{% endfor %}")
custom = out / "tpl_custom.docx"
doc.save(custom)
r3 = D.render(custom, ctx, out / "rendered_custom.docx")

from docx import Document as Doc2
txt = "\n".join(p.text for p in Doc2(str(r3)).paragraphs)
print("--- custom render ---")
print(txt)
assert "$12,500.50" in txt
assert "September 30, 2026" in txt
assert "Design, Copywriting" in txt
assert "Signed: Yes" in txt

btxt = "\n".join(c.text for t in Doc2(str(r2)).tables for r in t.rows for c in r.cells)
assert "Acme Rebrand" in btxt and "Beta Launch" in btxt, btxt
print("--- board table cells ---")
print(btxt)

# out-path templating
p = D._out_path("out/Quote - {{ name }} - {{ c.client_name }}.docx", ctx)
print("out path:", p)
assert p.name == "Quote - Acme Rebrand - Acme Corp.docx"
print("\nALL CHECKS PASSED")

itxt = "\n".join(c.text for t in Doc2(str(r1)).tables for r in t.rows for c in r.cells)
assert "Logo" in itxt and "Website" in itxt, "subitem loop dropped: " + itxt
assert "Acme Corp" in itxt and "$" not in itxt
print("item-template subitem loop OK")

# XML-special characters must survive rendering (monday values often contain "&")
doc = Document()
doc.add_paragraph("Category: {{ c.services | join_list }}")
doc.add_paragraph("Note: {{ c.client_name }}")
amp = out / "tpl_amp.docx"
doc.save(amp)
amp_item = {**item, "column_values": [
    cv("dropdown", "Services", "dropdown", "Meals & Entertainment"),
    cv("text", "Client Name", "text", "Smith & Sons <Legal> \"quoted\""),
]}
amp_ctx = D.base_context() | D.build_item_context(amp_item)
r4 = D.render(amp, amp_ctx, out / "rendered_amp.docx")
atxt = "\n".join(p.text for p in Doc2(str(r4)).paragraphs)
assert "Meals & Entertainment" in atxt, "ampersand lost: " + atxt
assert "Smith & Sons <Legal> \"quoted\"" in atxt, "special chars lost: " + atxt
print("XML-special characters preserved:", atxt.replace("\n", " | "))
